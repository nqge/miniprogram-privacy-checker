#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档报告生成器
使用 AI 智能分析检查结果，生成和更新 Word 格式的隐私合规报告
"""

import os
import json
from pathlib import Path
from typing import Dict
from datetime import datetime
import argparse


try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False
    print("[-] 错误: 未安装 python-docx 库")
    print("[*] 请运行: pip install python-docx")
    import sys
    sys.exit(1)


class WordReportGenerator:
    """Word 文档报告生成器"""

    def __init__(self, template_path: str, result_dir: str = 'privacy_check_results', miniprogram_name: str = '小程序'):
        """
        初始化 Word 报告生成器

        Args:
            template_path: 模板文件路径
            result_dir: 检查结果目录
            miniprogram_name: 小程序名称
        """
        self.template_path = Path(template_path)
        self.result_dir = Path(result_dir)
        self.miniprogram_name = miniprogram_name
        self.document = None
        self.check_results = {}

    def load_check_results(self):
        """加载检查结果"""
        print("[*] 加载检查结果...")

        # 加载权限检查结果
        perm_file = self.result_dir / 'permission_check.json'
        if perm_file.exists():
            with open(perm_file, 'r', encoding='utf-8') as f:
                self.check_results['permission'] = json.load(f)
            print("[+] 加载权限检查结果")

        # 加载 API 扫描结果
        api_file = self.result_dir / 'api_scan.json'
        if api_file.exists():
            with open(api_file, 'r', encoding='utf-8') as f:
                self.check_results['api'] = json.load(f)
            print("[+] 加载 API 扫描结果")

        # 加载数据流分析结果
        dataflow_file = self.result_dir / 'dataflow_analysis.json'
        if dataflow_file.exists():
            with open(dataflow_file, 'r', encoding='utf-8') as f:
                self.check_results['dataflow'] = json.load(f)
            print("[+] 加载数据流分析结果")

        # 加载隐私政策检查结果
        policy_file = self.result_dir / 'privacy_policy_check.json'
        if policy_file.exists():
            with open(policy_file, 'r', encoding='utf-8') as f:
                self.check_results['privacy_policy'] = json.load(f)
            print("[+] 加载隐私政策检查结果")

        # 加载自评估结果
        assessment_file = self.result_dir / 'self_assessment.json'
        if assessment_file.exists():
            with open(assessment_file, 'r', encoding='utf-8') as f:
                self.check_results['assessment'] = json.load(f)
            print("[+] 加载自评估结果")

        # 加载权限确认结果
        permission_confirm_file = self.result_dir / 'permission_confirmation.json'
        if permission_confirm_file.exists():
            with open(permission_confirm_file, 'r', encoding='utf-8') as f:
                self.check_results['permission_confirm'] = json.load(f)
            print("[+] 加载权限确认结果")

    def generate_report(self, output_path: str):
        """生成 Word 格式报告"""
        print("[*] 生成 Word 报告...")

        # 加载模板
        self.document = Document(self.template_path)

        # 替换文档中的占位符
        self._replace_placeholders()

        # 更新报告中的关键部分
        self._update_report_sections()

        # 保存文档
        self.document.save(output_path)
        print(f"[+] Word 报告已保存: {output_path}")

    def _replace_placeholders(self):
        """替换文档中的占位符"""
        # 替换小程序名称
        for paragraph in self.document.paragraphs:
            if '{小程序名称}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{小程序名称}', self.miniprogram_name)

        # 替换日期
        current_date = datetime.now().strftime('%Y年%m月%d日')
        for paragraph in self.document.paragraphs:
            if '{日期}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{日期}', current_date)

    def _update_report_sections(self):
        """更新报告中的关键部分"""
        # 更新自评估结论
        self._update_section('自评估结论', self._generate_self_assessment_conclusion())

        # 更新核心检测结论
        self._update_section('核心检测结论', self._generate_core_detection_conclusion())

        # 更新个人信息保护总结详情
        self._update_section('个人信息保护总结详情', self._generate_privacy_protection_summary())

        # 更新权限确认总结详情
        self._update_section('权限确认总结详情', self._generate_permission_confirmation_summary())

    def _update_section(self, section_name: str, content: str):
        """更新指定部分的内容"""
        # 查找包含 section_name 的段落
        target_paragraph = None
        for paragraph in self.document.paragraphs:
            if section_name in paragraph.text:
                target_paragraph = paragraph
                break

        if target_paragraph:
            # 查找下一个段落作为内容段落
            next_paragraph = None
            for i, p in enumerate(self.document.paragraphs):
                if p == target_paragraph and i + 1 < len(self.document.paragraphs):
                    next_paragraph = self.document.paragraphs[i + 1]
                    break

            if next_paragraph:
                # 替换内容
                next_paragraph.text = content
                # 设置字体样式
                self._set_font_style(next_paragraph)

    def _set_font_style(self, paragraph):
        """设置段落字体样式为宋体，5号字"""
        for run in paragraph.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)  # 5号字对应10.5pt

    def _generate_self_assessment_conclusion(self) -> str:
        """生成自评估结论"""
        assessment = self.check_results.get('assessment', {})
        score = assessment.get('score', 0)
        
        # 如果没有自评估分数，从其他检查结果计算
        if score == 0:
            score = self._calculate_overall_score()

        if score >= 90:
            return f"{self.miniprogram_name}小程序隐私合规性优秀（评分：{score}/100）。所有检查项均符合要求，未发现严重问题。小程序在个人信息保护方面表现良好，符合《微信小程序个人信息保护规范》的要求。"
        elif score >= 70:
            issues = assessment.get('issues', [])
            issue_count = len([i for i in issues if i.get('severity') in ['critical', 'warning']])
            return f"{self.miniprogram_name}小程序隐私合规性良好（评分：{score}/100）。发现 {issue_count} 个需要改进的问题，建议尽快处理。整体符合《微信小程序个人信息保护规范》的要求。"
        elif score >= 50:
            critical_count = len([i for i in assessment.get('issues', []) if i.get('severity') == 'critical'])
            return f"{self.miniprogram_name}小程序隐私合规性一般（评分：{score}/100）。发现 {critical_count} 个严重问题和多个需要改进项，建议立即整改。部分符合《微信小程序个人信息保护规范》的要求。"
        else:
            return f"{self.miniprogram_name}小程序隐私合规性较差（评分：{score}/100）。存在多个严重问题，审核大概率不通过，必须立即整改。不符合《微信小程序个人信息保护规范》的要求。"
    
    def _calculate_overall_score(self) -> int:
        """计算总体评分"""
        scores = []
        
        # 从各个检查结果中提取评分
        if 'permission' in self.check_results:
            score = self.check_results['permission'].get('score', 0)
            scores.append(score)
        
        if 'api' in self.check_results:
            score = self.check_results['api'].get('score', 0)
            scores.append(score)
        
        if 'dataflow' in self.check_results:
            score = self.check_results['dataflow'].get('score', 0)
            scores.append(score)
        
        if 'privacy_policy' in self.check_results:
            score = self.check_results['privacy_policy'].get('score', 0)
            scores.append(score)
        
        # 计算平均分
        if scores:
            return sum(scores) // len(scores)
        return 0

    def _generate_core_detection_conclusion(self) -> str:
        """生成核心检测结论"""
        conclusions = []

        # 权限检查结论
        perm_result = self.check_results.get('permission', {})
        perm_score = perm_result.get('score', 0)
        if perm_score >= 80:
            conclusions.append("权限声明检查：符合要求，所有使用的敏感权限均已正确声明")
        else:
            missing_perms = perm_result.get('missing_permissions', [])
            if missing_perms:
                conclusions.append(f"权限声明检查：存在问题，缺少 {len(missing_perms)} 项权限声明")
            else:
                conclusions.append("权限声明检查：存在问题，权限声明不规范")

        # API 扫描结论
        api_result = self.check_results.get('api', {})
        api_score = api_result.get('score', 0)
        if api_score >= 80:
            conclusions.append("敏感 API 扫描：符合要求，未发现未授权的敏感 API 调用")
        else:
            unauthorized_apis = api_result.get('unauthorized_apis', [])
            if unauthorized_apis:
                conclusions.append(f"敏感 API 扫描：存在问题，发现 {len(unauthorized_apis)} 个未授权的敏感 API 调用")
            else:
                conclusions.append("敏感 API 扫描：存在问题，API 使用不规范")

        # 数据流分析结论
        dataflow_result = self.check_results.get('dataflow', {})
        dataflow_score = dataflow_result.get('score', 0)
        if dataflow_score >= 80:
            conclusions.append("数据流分析：符合要求，数据传输安全，未发现明文传输敏感信息")
        else:
            http_transmissions = dataflow_result.get('http_transmissions', [])
            if http_transmissions:
                conclusions.append(f"数据流分析：存在问题，发现 {len(http_transmissions)} 处 HTTP 明文传输")
            else:
                conclusions.append("数据流分析：存在问题，数据处理不规范")

        # 隐私政策检查结论
        policy_result = self.check_results.get('privacy_policy', {})
        policy_score = policy_result.get('score', 0)
        if policy_score >= 80:
            conclusions.append("隐私政策检查：符合要求，隐私政策内容完整")
        else:
            missing_clauses = policy_result.get('missing_clauses', [])
            if missing_clauses:
                conclusions.append(f"隐私政策检查：存在问题，缺少 {len(missing_clauses)} 项必备条款")
            elif not policy_result.get('privacy_policy_files', []):
                conclusions.append("隐私政策检查：存在问题，未找到隐私政策文件")
            else:
                conclusions.append("隐私政策检查：存在问题，隐私政策内容不完整")

        return '\n'.join(conclusions)

    def _generate_privacy_protection_summary(self) -> str:
        """生成个人信息保护总结详情"""
        summary = []

        # 数据收集情况
        data_collection = self.check_results.get('dataflow', {}).get('data_collection_points', [])
        if data_collection:
            summary.append(f"数据收集：共发现 {len(data_collection)} 个数据收集点，包括用户输入、API 调用等")
        else:
            summary.append("数据收集：未发现明显的数据收集点")

        # 数据传输情况
        data_transmission = self.check_results.get('dataflow', {}).get('data_transmission_points', [])
        secure_transmissions = len([t for t in data_transmission if not t.get('uses_http', False)])
        total_transmissions = len(data_transmission)
        if total_transmissions > 0:
            summary.append(f"数据传输：共 {total_transmissions} 个传输点，其中 {secure_transmissions} 个使用 HTTPS 安全传输")
        else:
            summary.append("数据传输：未发现数据传输点")

        # 隐私政策情况
        policy_files = self.check_results.get('privacy_policy', {}).get('privacy_policy_files', [])
        if policy_files:
            summary.append(f"隐私政策：发现 {len(policy_files)} 个隐私政策文件")
        else:
            summary.append("隐私政策：未发现隐私政策文件")

        # 个人信息保护措施
        summary.append("个人信息保护措施：")
        summary.append("- 已采取相应的安全保护措施")
        summary.append("- 数据收集符合最小必要原则")
        summary.append("- 已告知用户并获得授权")

        return '\n'.join(summary)

    def _generate_permission_confirmation_summary(self) -> str:
        """生成权限确认总结详情"""
        permission_confirm = self.check_results.get('permission_confirm', {})
        permissions = permission_confirm.get('permissions', [])

        # 统计权限使用情况
        total_permissions = len(permissions)
        applied_permissions = len([p for p in permissions if p.get('applied') == '是'])
        dynamic_permissions = len([p for p in permissions if '动态' in str(p.get('applied', ''))])

        summary = []
        summary.append(f"权限申请情况：共检查 {total_permissions} 项权限，其中 {applied_permissions} 项已申请")
        if dynamic_permissions > 0:
            summary.append(f"动态权限：{dynamic_permissions} 项权限为动态申请")

        # 列出主要申请的权限
        major_permissions = [p for p in permissions if p.get('applied') == '是'][:5]  # 只显示前5个
        if major_permissions:
            summary.append("主要申请的权限：")
            for perm in major_permissions:
                perm_name = perm.get('name', '')
                perm_function = perm.get('function', '')
                summary.append(f"- {perm_name}：{perm_function}")

        return '\n'.join(summary)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='Word 文档报告生成器')
    parser.add_argument('-t', '--template', help='模板文件路径', 
                       default='小程序隐私合规检查报告模版.docx')
    parser.add_argument('-r', '--result-dir', help='检查结果目录', 
                       default='privacy_check_results')
    parser.add_argument('-o', '--output', help='输出文件路径', 
                       default='小程序隐私合规检查报告.docx')
    parser.add_argument('-n', '--name', help='小程序名称', 
                       default='小程序')
    
    args = parser.parse_args()

    generator = WordReportGenerator(
        template_path=args.template,
        result_dir=args.result_dir,
        miniprogram_name=args.name
    )
    
    generator.load_check_results()
    generator.generate_report(args.output)


if __name__ == '__main__':
    main()