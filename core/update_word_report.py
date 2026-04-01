#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新Word报告中的AI分析内容
"""

import os
import re
from docx import Document


def generate_personal_info_summary(assessment_file):
    """
    基于自评估表生成个人信息保护总结详情
    
    Args:
        assessment_file: 自评估表文件路径
    
    Returns:
        str: 个人信息保护总结详情
    """
    try:
        with open(assessment_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取总体评估结果
        score_match = re.search(r'评估得分：(\d+/\d+)', content)
        level_match = re.search(r'合规等级：(\w+)', content)
        
        score = score_match.group(1) if score_match else "未知"
        level = level_match.group(1) if level_match else "未知"
        
        # 提取评估说明
        assess_description = ""
        description_match = re.search(r'评估说明：\n(.*?)(?=\n- 改进建议：|$)', content, re.DOTALL)
        if description_match:
            assess_description = description_match.group(1).strip()
        
        # 提取改进建议
        suggestions = []
        suggestions_match = re.search(r'- 改进建议：\n(.*?)(?=\n================================================================================|$)', content, re.DOTALL)
        if suggestions_match:
            suggestions_text = suggestions_match.group(1).strip()
            suggestions = [line.strip('- ').strip() for line in suggestions_text.split('\n') if line.strip()]
        
        # 构建总结
        summary = f"个人信息保护自评估结果显示，小程序整体合规性{level}，评估得分为{score}。\n"
        summary += f"评估说明：{assess_description}\n\n"
        
        if suggestions:
            summary += "改进建议：\n"
            for i, suggestion in enumerate(suggestions, 1):
                summary += f"{i}. {suggestion}\n"
        
        return summary
    except Exception as e:
        print(f"生成个人信息保护总结失败: {e}")
        return "个人信息保护总结生成失败"


def generate_permission_summary(permission_file):
    """
    基于权限确认单生成权限确认总结详情
    
    Args:
        permission_file: 权限确认单文件路径
    
    Returns:
        str: 权限确认总结详情
    """
    try:
        with open(permission_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取申请的权限
        granted_permissions = []
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if '小程序是否申请：是' in line:
                # 向前查找权限名称
                permission_name = ""
                for j in range(i-5, i):
                    if j >= 0 and lines[j].strip() and not lines[j].strip().startswith('-'):
                        permission_name = lines[j].strip()
                        break
                # 向后查找业务功能和必要性说明
                business_function = ""
                necessity = ""
                for j in range(i+1, min(i+5, len(lines))):
                    if '对应业务功能：' in lines[j]:
                        business_function = lines[j].split('：', 1)[1].strip()
                    elif '必要性说明：' in lines[j]:
                        necessity = lines[j].split('：', 1)[1].strip()
                if permission_name:
                    granted_permissions.append({
                        "name": permission_name,
                        "business": business_function,
                        "necessity": necessity
                    })
        
        # 提取合规性评估
        compliance_assessment = ""
        compliance_match = re.search(r'权限声明合规性：(.*?)(?=\n注意事项:|$)', content, re.DOTALL)
        if compliance_match:
            compliance_assessment = compliance_match.group(1).strip()
        
        # 提取注意事项
        notes = []
        notes_match = re.search(r'注意事项：\n(.*?)(?=\n================================================================================|$)', content, re.DOTALL)
        if notes_match:
            notes_text = notes_match.group(1).strip()
            notes = [line.strip('   ').strip() for line in notes_text.split('\n') if line.strip()]
        
        # 构建总结
        summary = f"小程序共申请了{len(granted_permissions)}项权限，具体如下：\n"
        for i, perm in enumerate(granted_permissions, 1):
            summary += f"{i}. {perm['name']}\n"
            summary += f"   对应业务功能：{perm['business']}\n"
            summary += f"   必要性说明：{perm['necessity']}\n"
        
        summary += f"\n权限声明合规性：{compliance_assessment}\n\n"
        
        if notes:
            summary += "注意事项：\n"
            for i, note in enumerate(notes, 1):
                summary += f"{i}. {note}\n"
        
        return summary
    except Exception as e:
        print(f"生成权限确认总结失败: {e}")
        return "权限确认总结生成失败"


def update_word_report(docx_path, personal_info_summary, permission_summary):
    """
    更新Word报告中的AI分析内容
    
    Args:
        docx_path: Word文档路径
        personal_info_summary: 个人信息保护总结详情
        permission_summary: 权限确认总结详情
    
    Returns:
        bool: 更新是否成功
    """
    try:
        doc = Document(docx_path)
        updated_count = 0
        
        # 遍历所有段落，替换{AI更新内容}部分
        for i, paragraph in enumerate(doc.paragraphs):
            if "{AI更新内容}" in paragraph.text:
                # 检查前一个段落的内容来判断应该替换为什么内容
                if i > 0 and "个人信息保护总结详情" in doc.paragraphs[i-1].text:
                    paragraph.text = paragraph.text.replace("{AI更新内容}", personal_info_summary)
                    updated_count += 1
                    print(f"已更新段落 {i}: 个人信息保护总结详情")
                elif i > 0 and "权限确认总结详情" in doc.paragraphs[i-1].text:
                    paragraph.text = paragraph.text.replace("{AI更新内容}", permission_summary)
                    updated_count += 1
                    print(f"已更新段落 {i}: 权限确认总结详情")
                else:
                    print(f"段落 {i} 包含 {{AI更新内容}}，但无法确定替换内容类型")
                    print(f"前一个段落: {doc.paragraphs[i-1].text if i > 0 else '无'}")
        
        # 保存更新后的文档
        doc.save(docx_path)
        print(f"Word报告已更新: {docx_path}, 共更新 {updated_count} 个段落")
        return True
    except Exception as e:
        print(f"更新Word报告失败: {e}")
        return False


def main():
    """
    主函数
    """
    import argparse
    
    parser = argparse.ArgumentParser(description="更新Word报告中的AI分析内容")
    parser.add_argument("docx_path", type=str, help="Word文档路径")
    parser.add_argument("--assessment", type=str, default="自评估表.txt", help="自评估表文件路径")
    parser.add_argument("--permission", type=str, default="权限确认单.txt", help="权限确认单文件路径")
    args = parser.parse_args()
    
    # 生成总结
    personal_info_summary = generate_personal_info_summary(args.assessment)
    permission_summary = generate_permission_summary(args.permission)
    
    # 更新Word报告
    update_word_report(args.docx_path, personal_info_summary, permission_summary)


if __name__ == "__main__":
    main()
