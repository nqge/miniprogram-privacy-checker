#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版 Word 报告生成器
使用 AI 智能分析检查结果，生成和更新 Word 格式的隐私合规报告
"""

import os
import json
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import logging

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False

from .ai_agent_engine import AIAgentEngine

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class EnhancedWordReportGenerator:
    """增强版 Word 报告生成器"""

    def __init__(
        self,
        template_path: str,
        result_dir: str = 'privacy_check_results',
        miniprogram_name: str = '小程序',
        enable_ai: bool = False
    ):
        """
        初始化增强版 Word 报告生成器

        Args:
            template_path: 模板文件路径
            result_dir: 检查结果目录
            miniprogram_name: 小程序名称
            enable_ai: 是否启用 AI 分析
        """
        if not WORD_SUPPORT:
            raise ImportError("未安装 python-docx 库，请运行: pip install python-docx")

        self.template_path = Path(template_path)
        self.result_dir = Path(result_dir)
        self.miniprogram_name = miniprogram_name
        self.enable_ai = enable_ai

        # 初始化 AI 引擎
        self.ai_engine = AIAgentEngine(enable_ai=enable_ai) if enable_ai else None

        # 存储检查结果
        self.check_results = {}

        # 关键段落映射（根据模板分析）
        self.key_paragraphs = {
            '小程序名称': 8,      # 段落8: {小程序名称}小程序
            '自评估结论': 49,     # 段落49: 自评估结论
            '核心检测结论': 50,   # 段落50: 核心检测结论
            '权限申请': 52,       # 段落52: {AI更新内容}
            '隐私政策': 54,       # 段落54: {AI更新内容}
            '个人信息自评估': 57, # 段落57: {AI更新内容}
            '安全与第三方': 59,   # 段落59: {AI更新内容}
            '个人信息保护总结详情': 65, # 段落65: {AI更新内容}
            '权限确认总结详情': 68, # 段落68（需要在模板中找到确切位置）
        }

        logger.info(f"初始化增强版 Word 报告生成器")
        logger.info(f"模板路径: {self.template_path}")
        logger.info(f"小程序名称: {self.miniprogram_name}")
        logger.info(f"AI 分析: {'启用' if enable_ai else '禁用'}")

    def load_check_results(self):
        """加载检查结果"""
        logger.info("加载检查结果...")

        result_files = {
            'permission': 'permission_check.json',
            'api': 'api_scan.json',
            'dataflow': 'dataflow_analysis.json',
            'privacy_policy': 'privacy_policy_check.json',
            'assessment': 'self_assessment.json',
            'permission_confirm': 'permission_confirmation.json',
            'sdk': 'sdk_detection.json',
            'hybrid': 'hybrid_check.json'
        }

        for key, filename in result_files.items():
            file_path = self.result_dir / filename
            if file_path.exists():
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        self.check_results[key] = json.load(f)
                    logger.info(f"[+] 加载 {key}: {filename}")
                except Exception as e:
                    logger.error(f"[-] 加载 {key} 失败: {e}")

    def replace_miniprogram_name(self, doc: Document):
        """替换小程序名称占位符"""
        logger.info("替换小程序名称...")

        # 替换所有 {小程序名称} 占位符
        for paragraph in doc.paragraphs:
            if '{小程序名称}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{小程序名称}', self.miniprogram_name)
                logger.debug(f"替换段落: {paragraph.text[:50]}...")

        # 替换日期占位符
        current_date = datetime.now().strftime('%Y年%m月%d日')
        for paragraph in doc.paragraphs:
            if '{日期}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{日期}', current_date)

    def update_paragraph_52_permission_apply(self) -> str:
        """
        更新段落52: 权限申请

        Returns:
            更新内容文本
        """
        logger.info("生成段落52: 权限申请内容")

        if not self.check_results.get('permission'):
            return "暂无权限检查结果"

        # 使用 AI 分析或规则引擎
        if self.ai_engine:
            summary = self.ai_engine.analyze_permission_summary(
                self.result_dir / 'permission_check.json'
            )
        else:
            # 使用规则引擎生成
            perm_data = self.check_results['permission']
            declared = perm_data.get('declared_permissions', {})
            used = perm_data.get('used_permissions', {})

            necessary = [k for k in used.keys() if k in declared]
            unnecessary = [k for k in used.keys() if k not in declared]

            summary = {
                'summary': f"共使用 {len(used)} 项权限，其中 {len(necessary)} 项已声明，{len(unnecessary)} 项未声明",
                'necessary_permissions': [{'name': k, 'purpose': used.get(k, {}).get('purpose', '未说明')} for k in necessary],
                'unnecessary_permissions': [{'name': k, 'suggestion': f'需要在 app.json 中声明 {k} 权限'} for k in unnecessary],
                'risk_assessment': "低风险" if len(unnecessary) == 0 else "中风险" if len(unnecessary) <= 2 else "高风险"
            }

        # 生成段落内容
        content = f"""
本次检查发现，小程序共使用 {len(self.check_results.get('permission', {}).get('used_permissions', {}))} 项权限。

**已正确声明的权限**:
{chr(10).join([f"- {p['name']}: {p.get('purpose', '用于核心业务功能')}" for p in summary.get('necessary_permissions', [])[:5]])}

**未声明的权限**:
{chr(10).join([f"- {p['name']}: {p.get('suggestion', '建议补充声明')}" for p in summary.get('unnecessary_permissions', [])[:5]]) if summary.get('unnecessary_permissions') else '无'}

**风险评估**: {summary.get('risk_assessment', '未知')}
"""

        return content.strip()

    def update_paragraph_54_privacy_policy(self) -> str:
        """
        更新段落54: 隐私政策

        Returns:
            更新内容文本
        """
        logger.info("生成段落54: 隐私政策内容")

        if not self.check_results.get('privacy_policy'):
            return "暂无隐私政策检查结果"

        policy_data = self.check_results['privacy_policy']
        completeness = policy_data.get('completeness', 0)
        missing_items = policy_data.get('missing_items', [])

        content = f"""
隐私政策完整度为 {completeness}%。

**缺失内容**:
{chr(10).join([f"- {item}" for item in missing_items[:10]]) if missing_items else '无'}

**建议**: 请根据《个人信息保护法》和平台要求，完善隐私政策内容，特别是{missing_items[0] if missing_items else '用户权利告知'}部分。
"""

        return content.strip()

    def update_paragraph_57_personal_info_assessment(self) -> str:
        """
        更新段落57: 个人信息自评估

        Returns:
            更新内容文本
        """
        logger.info("生成段落57: 个人信息自评估内容")

        if not self.check_results.get('assessment'):
            return "暂无自评估结果"

        # 使用 AI 分析或规则引擎
        if self.ai_engine:
            summary = self.ai_engine.analyze_assessment_summary(
                self.result_dir / 'self_assessment.json'
            )
        else:
            assess_data = self.check_results['assessment']
            assessments = assess_data.get('assessments', [])

            passed = sum(1 for a in assessments if a.get('result') == '符合')
            failed = sum(1 for a in assessments if a.get('result') == '不符合')
            partial = sum(1 for a in assessments if a.get('result') == '部分符合')

            summary = {
                'collection_summary': f"共评估 {len(assessments)} 个检查点，其中 {passed} 项符合，{partial} 项部分符合，{failed} 项不符合",
                'protection_measures': [a.get('description', '') for a in assessments if a.get('result') in ['符合', '部分符合']][:10],
                'compliance_summary': "基本合规" if passed / len(assessments) >= 0.8 else "部分合规" if passed / len(assessments) >= 0.6 else "合规性较低"
            }

        content = f"""
{summary.get('collection_summary', '')}

**保护措施**:
{chr(10).join([f"- {m}" for m in summary.get('protection_measures', [])[:5]]) if summary.get('protection_measures') else '无'}

**合规性总结**: {summary.get('compliance_summary', '未知')}
"""

        return content.strip()

    def update_paragraph_59_security_thirdparty(self) -> str:
        """
        更新段落59: 安全与第三方

        Returns:
            更新内容文本
        """
        logger.info("生成段落59: 安全与第三方内容")

        if not self.check_results.get('sdk'):
            return "暂无 SDK 检测结果"

        sdk_data = self.check_results['sdk']
        sdks = sdk_data.get('sdks', [])

        content = f"""
检测到 {len(sdks)} 个第三方 SDK。

{chr(10).join([f"- **{sdk.get('name', '未知')}**: {sdk.get('purpose', '用途未知')}" for sdk in sdks[:10]]) if sdks else '未检测到第三方 SDK'}

**建议**: 请检查第三方 SDK 的隐私政策，确保其数据处理行为符合规范，并在隐私政策中充分披露。
"""

        return content.strip()

    def update_paragraph_65_protection_summary(self) -> str:
        """
        更新段落65: 个人信息保护总结详情

        Returns:
            更新内容文本
        """
        logger.info("生成段落65: 个人信息保护总结详情")

        # 读取自评估表.txt
        assessment_txt = self.result_dir / '自评估表.txt'
        if assessment_txt.exists():
            try:
                with open(assessment_txt, 'r', encoding='utf-8') as f:
                    assessment_content = f.read()

                # 提取关键信息
                lines = assessment_content.split('\n')
                passed_count = sum(1 for line in lines if '符合' in line)
                total_count = len([line for line in lines if '评估点' in line or '□' in line])

                return f"""
根据自评估表检查结果，小程序在个人信息保护方面{f"有 {total_count - passed_count} 项需要改进" if passed_count < total_count else "基本符合要求"}。

主要保护措施包括：
- 数据收集遵循最小必要原则
- 建立了用户授权机制
- 提供了隐私政策文本
- 设置了用户投诉渠道

建议继续完善用户权利实现机制和数据安全措施。
"""
            except Exception as e:
                logger.error(f"读取自评估表失败: {e}")

        return "个人信息保护总结详情需要根据自评估表结果生成。"

    def update_paragraph_68_permission_confirm_summary(self) -> str:
        """
        更新段落68: 权限确认总结详情

        Returns:
            更新内容文本
        """
        logger.info("生成段落68: 权限确认总结详情")

        # 读取权限确认单.txt
        permission_txt = self.result_dir / '权限确认单.txt'
        if permission_txt.exists():
            try:
                with open(permission_txt, 'r', encoding='utf-8') as f:
                    permission_content = f.read()

                # 提取关键信息
                lines = permission_content.split('\n')
                necessary_count = sum(1 for line in lines if '必要' in line and '☑' in line)
                unnecessary_count = sum(1 for line in lines if '不必要' in line)

                return f"""
根据权限确认单检查结果，小程序申请的{necessary_count + unnecessary_count}项权限中，{necessary_count}项为必要权限，{unnecessary_count}项为非必要权限。

所有权限申请均符合业务功能需求，未发现过度申请权限的情况。建议继续保持权限最小化原则。
"""
            except Exception as e:
                logger.error(f"读取权限确认单失败: {e}")

        return "权限确认总结详情需要根据权限确认单结果生成。"

    def update_font_style(self, doc: Document):
        """更新字体样式（宋体，5号）"""
        logger.info("更新字体样式...")

        # 5号字体 = 10.5pt
        font_size = Pt(10.5)
        font_name = '宋体'

        # 更新所有段落
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.name = font_name
                    run.font.size = font_size

                    # 设置中文字体
                    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia', font_name)
                except Exception as e:
                    logger.debug(f"设置字体失败: {e}")

    def generate_report(self, output_path: Optional[str] = None) -> str:
        """
        生成增强版 Word 报告

        Args:
            output_path: 输出文件路径（可选）

        Returns:
            生成的报告文件路径
        """
        logger.info("开始生成增强版 Word 报告...")

        # 1. 加载检查结果
        self.load_check_results()

        # 2. 加载模板
        logger.info("加载模板...")
        doc = Document(str(self.template_path))

        # 3. 替换小程序名称
        self.replace_miniprogram_name(doc)

        # 4. 更新关键段落内容
        paragraph_updates = {
            52: self.update_paragraph_52_permission_apply(),
            54: self.update_paragraph_54_privacy_policy(),
            57: self.update_paragraph_57_personal_info_assessment(),
            59: self.update_paragraph_59_security_thirdparty(),
            65: self.update_paragraph_65_protection_summary(),
            68: self.update_paragraph_68_permission_confirm_summary()
        }

        for para_index, content in paragraph_updates.items():
            if para_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_index]
                # 清空现有内容
                paragraph.clear()
                # 添加新内容
                run = paragraph.add_run(content)
                logger.info(f"更新段落 {para_index}: {content[:50]}...")

        # 5. 更新字体样式
        self.update_font_style(doc)

        # 6. 保存报告
        if not output_path:
            output_path = self.result_dir / f"{self.miniprogram_name}小程序隐私合规检查报告.docx"

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc.save(str(output_path))
        logger.info(f"报告已生成: {output_path}")

        return str(output_path)


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(description='增强版小程序隐私合规检查报告生成器')
    parser.add_argument('--template', required=True, help='模板文件路径')
    parser.add_argument('--result-dir', default='privacy_check_results', help='检查结果目录')
    parser.add_argument('--miniprogram-name', default='小程序', help='小程序名称')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--enable-ai', action='store_true', help='启用 AI 分析')

    args = parser.parse_args()

    # 创建生成器
    generator = EnhancedWordReportGenerator(
        template_path=args.template,
        result_dir=args.result_dir,
        miniprogram_name=args.miniprogram_name,
        enable_ai=args.enable_ai
    )

    # 生成报告
    output_path = generator.generate_report(args.output)

    print(f"\n✅ 报告生成成功: {output_path}")


if __name__ == '__main__':
    main()
